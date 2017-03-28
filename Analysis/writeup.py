# writeup.py
# by James Fulford

import docx  # used to write Word documents
import matplotlib.pyplot as plt  # used to make plots
import analytics as a  # personal package for statistics
from datetime import datetime as dt
from analytics.dataset import Dataset as DS  # personal tool for managing data
import string
import seaborn as sns

def remove_control_characters(s):
    filtered_string = "".join(filter(lambda x: x in string.printable, s))
    return filtered_string


def read(section):
    import string
    with open(section) as phile:
        return remove_control_characters(phile.read())


ds = DS.load("Confirmation Program.json")


doc = docx.Document()
doc.add_heading("Analysis of Souls", 0)
doc.add_heading("A Statistical Study of Confirmation Students")
doc.add_heading("By James Fulford\tFall 2016 - Spring 2017", 2)


# INTRODUCTION
data = {
    'n': a.count(ds.data)
}
intro = read("introduction.txt").format(**data)
doc.add_paragraph(intro)


# VARIABLE EXPLANATIONS
def cnt(filt_criteria, shr=False):
    lst = filter(filt_criteria, ds.data)
    ret = len(lst)
    if shr:
        ret = round(100 * ret / float(len(ds.data)), 1)
    return ret

variables = {
    "Confirmation": {
        "confirmed": cnt(lambda x: x["attendance"] is True),
        "not_confirmed": cnt(lambda x: x["attendance"] is False),
        "confirmed_share": cnt(lambda x: x["attendance"] is True, shr=True),
        "not_confirmed_share": cnt(lambda x: x["attendance"] is False, shr=True)
    },
    "Absenteeism": {
        "summer": cnt(lambda x: "Summer" in x["section"] and "Year" not in x["section"]),
        "summer_share": cnt(lambda x: "Summer" in x["section"] and "Year" not in x["section"], shr=True),
        "year": cnt(lambda x: "Summer" not in x["section"] and "Year" in x["section"]),
        "year_share": cnt(lambda x: "Summer" not in x["section"] and "Year" in x["section"], shr=True),
    },
    "Volunteering": {
        "no_vol": cnt(lambda x: len(x["volunteering"]) is 0),
        "not_confirmed": cnt(lambda x: x["attendance"] is False),
        "non_conf_no_vol": cnt(lambda x: len(x["volunteering"]) == 0 and x["attendance"] is False),
        "non_conf_with_vol": cnt(lambda x: x["attendance"] is False) - cnt(lambda x: len(x["volunteering"]) == 0 and x["attendance"] is False)
    }
}

for var in variables.keys():
    content = read(var.lower() + ".txt").format(**variables[var])
    doc.add_heading(var, 2)
    doc.add_paragraph(content)


# COHORT ENUMERATION

base = ds.data
cohorts = [
    ("Base", "Attended program within year of Confirmation ceremony", lambda x: True),
    ("Summer", "In Summer program", lambda x: "Summer" in x["section"] and "Year" not in x["section"]),
    ("Year", "In Year-round program", lambda x: "Year" in x["section"] and "Summer" not in x["section"]),
    ("Confirmed", "Confirmed year of class completion at St. Elizabeth Seton", lambda x: x["attendance"] is True),
    ("Not Confirmed", "Not Confirmed year of class completion or not at St. Elizabeth Seton", lambda x: x["attendance"] is False),
    ("No Vol.", "Did not have a volunteering record", lambda x: len(x["volunteering"]) is 0),
    ("Low Vol.", "Number of volunteering events below 5-event guideline", lambda x: len(x["volunteering"]) < 5 and len(x["volunteering"]) != 0),
    ("High Vol.", "Number of volunteering events above 8", lambda x: len(x["volunteering"]) > 8),
    ("2013", "Finished classes in 2013", lambda x: x["section"].__contains__("2013")),
    ("2014", "Finished classes in 2014", lambda x: x["section"].__contains__("2014")),
    ("2015", "Finished classes in 2015", lambda x: x["section"].__contains__("2015")),
    ("2016", "Finished classes in 2016", lambda x: x["section"].__contains__("2016")),
]

doc.add_heading("Cohorts", 2)
doc.add_paragraph(read("cohorts.txt").format(count_cohorts=len(cohorts)))
table = doc.add_table(rows=1, cols=3)
table.style = 'Medium List 2 Accent 1'
cols = table.columns
cols[0].width = docx.shared.Inches(1.07)
cols[1].width = docx.shared.Inches(4.38)
cols[2].width = docx.shared.Inches(0.62)

# Headers
table.cell(0, 0).text = "Cohort Title"
table.cell(0, 1).text = "Criteria"
table.cell(0, 2).text = "Count"

# Content
for i in range(len(cohorts)):
    cells = table.add_row().cells
    cells[0].text = cohorts[i][0]  # title
    cells[1].text = cohorts[i][1]  # description
    cells[2].text = str(cnt(cohorts[i][2]))  # lambda function counts


# VARIABLE DISTRIBUTIONS AND ONE-VARIABLE STATISTICS

doc.add_heading("Centers of Distributions", 1)

acs = {
    "Confirmation": lambda x: 1 if x["attendance"] is True else 0,
    "Absenteeism": lambda x: sum(x["absences"]),
    "Volunteering": lambda x: len(x["volunteering"]),
}

confidence = .95
# s presumed to be max-min/4 by the Empirical Rule


def get_data(var, cohort=None):
    if cohort:
        data = map(acs[var], filter(cohort[2], ds.data))
    else:
        data = map(acs[var], ds.data)
    return data

def get_one_var_stats(var, cohort=None):
    data = get_data(var, cohort=cohort)
    stats = {
        "variable": var,
        "mean": round(a.mean(data), 3),
        "datarange": a.datarange(data),
        "cfd_int": map(lambda x: round(x, 2), a.confidence_interval(data, conf=confidence)),
        "five-point-summary": a.five_point_summary(data),
        "confidence": confidence
    }
    # if cohort:
    #     # pass
    #     print var, cohort[0], round(a.mean(data), 2), round(a.standard_deviation(data), 2), stats["cfd_int"]
    return data, stats

for var in variables:
    data, stats = get_one_var_stats(var)

    # Writing in document:
    doc.add_heading(var, 2)
    doc.add_paragraph(read("one-var-stats.txt").format(**stats))

    for cohort in cohorts:
        try:
            co_data, co_stats = get_one_var_stats(var, cohort=cohort)
        except AssertionError:
            pass  # one of the calculations couldn't be made.
        except ZeroDivisionError:
            pass

        if co_stats["cfd_int"][0] > stats["mean"] or co_stats["cfd_int"][1] < stats["mean"]:
            # if cohort inference interval doesn't include total mean:
            doc.add_heading(var + ": " + cohort[0], 3)
            doc.add_paragraph(read("one-var-stats.txt").format(**co_stats))

    # SPECIAL CASES:
    if var in ["Volunteering"]:
        plt.hist(data, stats["datarange"])
        plt.axis([a.minimum(data), a.maximum(data), 0, a.maximum(a.frequency(data).values()) * 1.1])
        plt.title(var + " (classes: " + str(stats["datarange"]) + ")")
        plt.xlabel(var)
        plt.ylabel("Frequency (n: " + str(a.count(data)) + ")")
        plt.savefig(var + " distribution.png")
        doc.add_picture(var + " distribution.png", width=docx.shared.Inches(5), height=docx.shared.Inches(3))
        # plt.show()


# VARIABLE RELATIONSHIPS
doc.add_heading("Variable Relationships", 1)

def get_two_var_stats(var1, var2, cohort=None):
    data1 = get_data(var1, cohort=cohort)
    data2 = get_data(var2, cohort=cohort)

    r = a.pearson_r(data1, data2)
    if r is None:
        return data1, data2, {}  # just stop calculations now.

    stats = {
        "r": round(r, 3),
        "r2": round(r ** 2, 3),
        "explain-share": 100 * round(r ** 2, 3),
        "var1": var1,
        "var2": var2
    }
    if cohort:
        stats["cohort"] = cohort[0]
        # print var1, var2, cohort[0], stats["r"], stats['r2'], str(stats["explain-share"]) + "%"
        if cohort[0] == "Base":
            plt.clf()
            datum = {}
            datum[var1] = data1
            datum[var2] = data2
            # sns.lmplot(x=var1, y=var2, data=datum)
            # sns.despine()
            # plt.show()
    return data1, data2, stats

# Go through all relationships
keys = variables.keys()
datas = map(lambda var: map(acs[var], ds.data), keys)
for i in range(len(keys)):
    var1 = keys[i]
    for j in range(i + 1, len(keys)):
        var2 = keys[j]

        # COHORT ANALYSIS (filter data)
        for cohort in cohorts:
            try:
                data1, data2, co_stats = get_two_var_stats(var1, var2, cohort)
            except ZeroDivisionError:
                pass  # could not calculate
            if "r" in co_stats.keys() and abs(co_stats["r"]) > .3:
                doc.add_heading(var1 + " and " + var2 + ": " + cohort[0], 2)
                doc.add_paragraph(read("variable-relationship.txt").format(**co_stats))


# HYPOTHESIS TESTING
keys = variables.keys()
datas = map(lambda var: map(acs[var], ds.data), keys)


for i in range(len(keys)):
    var = keys[i]
    doc.add_heading(var, 2)
    table = doc.add_table(1, 5)
    table.style = 'Medium List 2 Accent 1'
    heading_cells = table.rows[0].cells
    heading_cells[0].text = 'Lower Cohort'
    heading_cells[1].text = 'Lower Mean'
    heading_cells[2].text = 'Greater Mean'
    heading_cells[3].text = 'Greater Cohort'
    heading_cells[4].text = 'p value'

    # Loop through cohorts
    for co_i in range(len(cohorts)):
        co1 = cohorts[co_i]
        # Loop through remaining cohorts
        for co_j in range(co_i + 1, len(cohorts)):
            co2 = cohorts[co_j]

            cohort1 = get_data(var, cohort=co1)
            cohort2 = get_data(var, cohort=co2)

            # INTERCOHORT ANALYSIS
            from analytics import statistics_and_tests as stts
            tolerance = .00000000000001
            try:
                assert a.standard_deviation(cohort1) > tolerance
                assert a.standard_deviation(cohort2) > tolerance
            except AssertionError:
                continue
            except ZeroDivisionError:
                continue
            # Ho: Mean of cohorts are equal
            # Ha: Means of cohorts are unequal
            too_alpha = 0.
            alpha = 0.05
            pval = stts.ds_likely_same_means(cohort1, cohort2, two_sided=False)
            if  pval < alpha and pval > too_alpha:
                # strong evidence against null hypothesis:
                # reject the null hypothesis with alpha confidence
                cells = table.add_row().cells
                if round(a.mean(cohort1), 3) < round(a.mean(cohort2), 3):
                    cells[0].text = co1[0]
                    cells[1].text = str(round(a.mean(cohort1), 3))
                    cells[2].text = str(round(a.mean(cohort2), 3))
                    cells[3].text = co2[0]
                    cells[4].text = str(round(pval, 3))
                    print var, co1[0], round(a.mean(cohort1), 3), round(a.mean(cohort2), 3), co2[0], str(round(100 * pval, 1)) + "%"
                else:
                    cells[0].text = co2[0]
                    cells[1].text = str(round(a.mean(cohort2), 3))
                    cells[2].text = str(round(a.mean(cohort1), 3))
                    cells[3].text = co1[0]
                    cells[4].text = str(round(pval, 3))
                    print var, co1[0], round(a.mean(cohort1), 3), round(a.mean(cohort2), 3), co2[0], str(round(100 * pval, 1)) + "%"



# Add in references
par = doc.paragraphs
runs = par[len(par) - 1].runs
runs[len(runs) - 1].add_break()

doc.add_heading("References", 2)
doc.add_paragraph(read("references.txt"))

doc.save("Analysis of Souls " + dt.now().strftime("%m-%d-%Y") + " James Fulford.docx")
